from __future__ import annotations

import re
from pathlib import Path
from uuid import uuid4

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from apps.api.app.core.config import get_settings
from apps.api.app.ports.interfaces import DocumentRenderPort
from apps.api.app.schemas.trademark import ApplicationDraftRequest


class RealDocumentRenderAdapter(DocumentRenderPort):
    port_name = "documentRender"
    provider_name = "docx-reportlab"
    mode = "real"

    def __init__(self) -> None:
        self.settings = get_settings()

    def availability(self) -> tuple[bool, str | None]:
        return True, None

    def render_application(self, payload: ApplicationDraftRequest, summary: dict, trace_id: str) -> tuple[str, str]:
        safe_name = re.sub(r"[^\w\u4e00-\u9fff\-]", "_", payload.trademark_name)[:40]
        base_name = f"{safe_name}-{uuid4().hex[:8]}"
        docx_path = Path(self.settings.generated_dir) / f"{base_name}.docx"
        pdf_path = Path(self.settings.generated_dir) / f"{base_name}.pdf"

        document = Document()
        document.add_heading("A1+ Trademark Application Draft", level=0)
        document.add_paragraph(f"Trademark Name: {payload.trademark_name}")
        document.add_paragraph(f"Applicant: {payload.applicant_name}")
        document.add_paragraph(f"Applicant Type: {payload.applicant_type}")
        document.add_paragraph(f"Suggested Classes: {', '.join(payload.categories)}")
        document.add_paragraph(f"Risk Level: {payload.risk_level}")
        document.add_paragraph(summary["summary"])
        document.add_paragraph(
            "Compliance Notice: A1+ only prepares documents and submission guidance. Users submit to official systems on their own."
        )
        document.save(docx_path)

        pdf = canvas.Canvas(str(pdf_path), pagesize=A4)
        pdf.drawString(72, 800, "A1+ Trademark Application Draft")
        pdf.drawString(72, 780, f"Trademark Name: {payload.trademark_name}")
        pdf.drawString(72, 760, f"Applicant: {payload.applicant_name}")
        pdf.drawString(72, 740, f"Classes: {', '.join(payload.categories)}")
        pdf.drawString(72, 720, f"Risk Level: {payload.risk_level}")
        pdf.drawString(72, 700, "A1+ only prepares files and submission guidance.")
        pdf.save()

        return str(docx_path), str(pdf_path)
