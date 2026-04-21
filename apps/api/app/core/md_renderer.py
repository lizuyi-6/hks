"""Unified Markdown → DOCX / PDF rendering pipeline.

All Word/PDF document generation in the backend is expected to first produce
a Markdown source of truth and then transform it through the functions in
this module. This keeps a single human-readable artifact (the Markdown) as
the canonical content while allowing downloaders to pick the delivery
format.

Design constraints
==================

- Pure Python. Only depends on ``python-docx`` and ``reportlab`` which are
  already pinned in :file:`apps/api/requirements.txt`. No system binaries
  such as ``pandoc`` or ``wkhtmltopdf`` are required so Docker images and CI
  don't change.
- Minimal parser. We deliberately support only the subset of Markdown that
  the product actually emits:

    * ATX headings (``#``, ``##``, ``###`` – up to level 3)
    * Paragraphs separated by blank lines
    * Unordered list items starting with ``- `` or ``* ``
    * Ordered list items starting with ``1. ``, ``2. `` …
    * Blockquotes starting with ``> ``
    * Horizontal rules (``---``)
    * Inline bold via ``**x**`` and italic via ``*x*``
    * Simple inline code via ```` `x` ````

  The two converters share the same AST so DOCX and PDF come out visually
  consistent.

The module purposefully does *not* try to be a general Markdown renderer –
swapping in ``pypandoc`` later is a single-file change without touching any
caller.
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Literal

logger = logging.getLogger(__name__)

BlockKind = Literal[
    "heading",
    "paragraph",
    "bullet",
    "ordered",
    "blockquote",
    "hr",
    "blank",
]


@dataclass
class Block:
    kind: BlockKind
    text: str = ""
    level: int = 0  # Heading level (1-3) or list ordinal when relevant
    items: list[str] = field(default_factory=list)  # For list blocks


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*\S)\s*$")
_ORDERED_RE = re.compile(r"^\s*(\d+)\.\s+(.*\S)\s*$")
_HR_RE = re.compile(r"^\s*(?:-{3,}|\*{3,}|_{3,})\s*$")
_BLOCKQUOTE_RE = re.compile(r"^\s*>\s?(.*)$")


def _parse(md: str) -> list[Block]:
    """Tokenise Markdown into a flat list of block-level elements.

    Lists are accumulated greedily until a non-list line breaks the run, so
    the downstream renderers can emit a proper list group rather than one
    item per block.
    """
    blocks: list[Block] = []
    lines = md.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if not line.strip():
            blocks.append(Block(kind="blank"))
            i += 1
            continue

        if _HR_RE.match(line):
            blocks.append(Block(kind="hr"))
            i += 1
            continue

        m = _HEADING_RE.match(line)
        if m:
            hashes, text = m.group(1), m.group(2).strip()
            blocks.append(Block(kind="heading", text=text, level=min(len(hashes), 3)))
            i += 1
            continue

        m = _BULLET_RE.match(line)
        if m:
            items = [m.group(1).strip()]
            i += 1
            while i < len(lines):
                m2 = _BULLET_RE.match(lines[i])
                if not m2:
                    break
                items.append(m2.group(1).strip())
                i += 1
            blocks.append(Block(kind="bullet", items=items))
            continue

        m = _ORDERED_RE.match(line)
        if m:
            items = [m.group(2).strip()]
            i += 1
            while i < len(lines):
                m2 = _ORDERED_RE.match(lines[i])
                if not m2:
                    break
                items.append(m2.group(2).strip())
                i += 1
            blocks.append(Block(kind="ordered", items=items))
            continue

        m = _BLOCKQUOTE_RE.match(line)
        if m:
            parts = [m.group(1).strip()]
            i += 1
            while i < len(lines):
                m2 = _BLOCKQUOTE_RE.match(lines[i])
                if not m2:
                    break
                parts.append(m2.group(1).strip())
                i += 1
            blocks.append(Block(kind="blockquote", text="\n".join(parts).strip()))
            continue

        # Default: collect consecutive non-special lines as a paragraph.
        para = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if (
                _HEADING_RE.match(nxt)
                or _BULLET_RE.match(nxt)
                or _ORDERED_RE.match(nxt)
                or _HR_RE.match(nxt)
                or _BLOCKQUOTE_RE.match(nxt)
            ):
                break
            para.append(nxt.strip())
            i += 1
        blocks.append(Block(kind="paragraph", text=" ".join(para).strip()))

    # Collapse trailing blanks; keep them only as separators between content.
    while blocks and blocks[-1].kind == "blank":
        blocks.pop()
    return blocks


# ---------------------------------------------------------------------------
# Inline text splitting (shared)
# ---------------------------------------------------------------------------


_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)",
)


@dataclass
class InlineRun:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


def _split_inline(text: str) -> list[InlineRun]:
    """Split a paragraph into bold / italic / inline-code runs."""
    if not text:
        return [InlineRun(text="")]
    parts = _INLINE_RE.split(text)
    runs: list[InlineRun] = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            runs.append(InlineRun(text=part[2:-2], bold=True))
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            runs.append(InlineRun(text=part[1:-1], italic=True))
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            runs.append(InlineRun(text=part[1:-1], code=True))
        else:
            runs.append(InlineRun(text=part))
    return runs


# ---------------------------------------------------------------------------
# DOCX writer
# ---------------------------------------------------------------------------


def md_to_docx(md: str, out_path: Path | str) -> Path:
    """Render Markdown to a ``.docx`` file via python-docx.

    Returns the absolute path that was written.
    """
    from docx import Document
    from docx.shared import Pt

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    # A sensible Chinese-friendly default; falls through to the system font
    # chain if the document is rendered on a host without "微软雅黑".
    try:
        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(11)
    except Exception:  # pragma: no cover - defensive, python-docx oddity
        logger.debug("failed to set Normal style font", exc_info=True)

    def add_inline(paragraph, runs: Iterable[InlineRun]) -> None:
        empty = True
        for r in runs:
            if not r.text:
                continue
            empty = False
            run = paragraph.add_run(r.text)
            run.bold = r.bold
            run.italic = r.italic
            if r.code:
                try:
                    run.font.name = "Consolas"
                except Exception:  # pragma: no cover
                    pass
        if empty:
            paragraph.add_run("")

    for block in _parse(md):
        if block.kind == "blank":
            continue
        if block.kind == "heading":
            doc.add_heading(block.text, level=block.level)
        elif block.kind == "paragraph":
            p = doc.add_paragraph()
            add_inline(p, _split_inline(block.text))
        elif block.kind == "bullet":
            for item in block.items:
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, _split_inline(item))
        elif block.kind == "ordered":
            for item in block.items:
                p = doc.add_paragraph(style="List Number")
                add_inline(p, _split_inline(item))
        elif block.kind == "blockquote":
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, _split_inline(block.text))
        elif block.kind == "hr":
            # python-docx has no true HR; add a thin separator paragraph.
            p = doc.add_paragraph("─" * 30)
            p.alignment = 1  # center

    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# PDF writer
# ---------------------------------------------------------------------------


_PDF_FONT_REGISTERED = False


def _register_pdf_fonts() -> str:
    """Register an embedded Chinese-capable font and return its name.

    Uses reportlab's bundled Adobe CID ``STSong-Light`` so Chinese text
    renders without needing a system TTF. Falls back to Helvetica on any
    registration failure (pure ASCII documents still look OK).
    """
    global _PDF_FONT_REGISTERED
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont

        if not _PDF_FONT_REGISTERED:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            _PDF_FONT_REGISTERED = True
        return "STSong-Light"
    except Exception:  # pragma: no cover - reportlab always has this font
        logger.exception("failed to register CID font; falling back to Helvetica")
        return "Helvetica"


def _inline_to_pdf_markup(text: str) -> str:
    """Translate a limited inline markdown subset into reportlab's
    Paragraph mini-HTML so bold/italic/code render visibly."""
    from xml.sax.saxutils import escape

    def repl(match: re.Match[str]) -> str:
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            return f"<b>{escape(token[2:-2])}</b>"
        if token.startswith("*") and token.endswith("*"):
            return f"<i>{escape(token[1:-1])}</i>"
        if token.startswith("`") and token.endswith("`"):
            return f'<font face="Courier">{escape(token[1:-1])}</font>'
        return escape(token)

    out: list[str] = []
    last = 0
    for m in _INLINE_RE.finditer(text):
        out.append(escape(text[last : m.start()]))
        out.append(repl(m))
        last = m.end()
    out.append(escape(text[last:]))
    return "".join(out)


def md_to_pdf(md: str, out_path: Path | str) -> Path:
    """Render Markdown to a ``.pdf`` file via reportlab.platypus."""
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    font_name = _register_pdf_fonts()
    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=base["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=16,
        alignment=TA_LEFT,
    )
    h_styles = {
        1: ParagraphStyle(
            "H1",
            parent=base["Heading1"],
            fontName=font_name,
            fontSize=18,
            leading=24,
            spaceBefore=12,
            spaceAfter=8,
        ),
        2: ParagraphStyle(
            "H2",
            parent=base["Heading2"],
            fontName=font_name,
            fontSize=14,
            leading=20,
            spaceBefore=10,
            spaceAfter=6,
        ),
        3: ParagraphStyle(
            "H3",
            parent=base["Heading3"],
            fontName=font_name,
            fontSize=12,
            leading=18,
            spaceBefore=8,
            spaceAfter=4,
        ),
    }
    quote = ParagraphStyle(
        "Quote",
        parent=body,
        leftIndent=18,
        textColor="#555",
        fontName=font_name,
    )

    doc = SimpleDocTemplate(
        str(out),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=out.stem,
    )

    story: list = []
    for block in _parse(md):
        if block.kind == "blank":
            story.append(Spacer(1, 6))
            continue
        if block.kind == "heading":
            style = h_styles.get(block.level, body)
            story.append(Paragraph(_inline_to_pdf_markup(block.text), style))
        elif block.kind == "paragraph":
            story.append(Paragraph(_inline_to_pdf_markup(block.text), body))
        elif block.kind in ("bullet", "ordered"):
            flow = ListFlowable(
                [
                    ListItem(
                        Paragraph(_inline_to_pdf_markup(item), body),
                        leftIndent=12,
                    )
                    for item in block.items
                ],
                bulletType="bullet" if block.kind == "bullet" else "1",
                bulletFontName=font_name,
                leftIndent=18,
            )
            story.append(flow)
        elif block.kind == "blockquote":
            story.append(Paragraph(_inline_to_pdf_markup(block.text), quote))
        elif block.kind == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5, color="#bbb"))

    if not story:
        # SimpleDocTemplate refuses to build an empty document.
        story.append(Paragraph("", body))

    doc.build(story)
    return out


# ---------------------------------------------------------------------------
# Convenience
# ---------------------------------------------------------------------------


def render_both(md: str, base_path: Path | str) -> tuple[Path, Path, Path]:
    """Write ``base_path`` + ``.md`` / ``.docx`` / ``.pdf`` and return the trio.

    The Markdown source is always persisted alongside its derived formats so
    debugging / auditing can compare what was produced.
    """
    base = Path(base_path)
    md_path = base.with_suffix(".md")
    docx_path = base.with_suffix(".docx")
    pdf_path = base.with_suffix(".pdf")
    md_path.parent.mkdir(parents=True, exist_ok=True)
    md_path.write_text(md, encoding="utf-8")
    md_to_docx(md, docx_path)
    md_to_pdf(md, pdf_path)
    return md_path, docx_path, pdf_path
